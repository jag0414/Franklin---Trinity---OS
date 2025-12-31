# Franklin OS • BidNova • Trinity
# Unified Resurrection Build

from __future__ import annotations

import asyncio
import json
import os
import uuid
from collections import deque
from datetime import datetime, timedelta, timezone
from typing import Any, Deque, Dict, List, Literal, Optional

import httpx
import jwt
from fastapi import FastAPI, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse
from pydantic import BaseModel, Field as PydanticField
from sqlalchemy import text
from sqlmodel import Field, Session, SQLModel, create_engine, select

# ---------------- CONFIG ----------------
APP_NAME = "Franklin OS • BidNova • Trinity"

SECRET = os.getenv("FRANKLIN_JWT_SECRET", "CHANGE_ME_NOW")
JWT_ALGO = "HS256"
DB_URL = os.getenv("FRANKLIN_DB_URL", "sqlite:///franklin.db")

# EMERGENCY FIX FOR RAILWAY/POSTGRES COMPATIBILITY
if DB_URL.startswith("postgres://"):
    DB_URL = DB_URL.replace("postgres://", "postgresql://", 1)

print(f"🚀 Franklin Engine: Connecting to {DB_URL.split('@')[-1]}...") # Safe logging

try:
    engine = create_engine(DB_URL, pool_pre_ping=True)
    with Session(engine) as session:
        # Force a simple query to verify connection
        session.execute(text("SELECT 1"))
    print("✅ Database Connection: SECURE")
except Exception as e:
    print(f"❌ Database Connection FAILED: {e}")
    # Fallback to local if production db fails (Sovereign Safety)
    engine = create_engine("sqlite:///franklin.db")

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY")
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
STABILITY_API_KEY = os.getenv("STABILITY_API_KEY")


def is_key_valid(key: Optional[str]) -> bool:
    if not key:
        return False
    if "..." in key or "CHANGE_ME" in key:
        return False
    if len(key) < 10:
        return False
    return True

# Bounded concurrency (safe defaults)
GLOBAL_MAX_CONCURRENCY = int(os.getenv("FRANKLIN_AI_GLOBAL_CONCURRENCY", "8"))
DEFAULT_PROVIDER_CONCURRENCY = int(os.getenv("FRANKLIN_PROVIDER_CONCURRENCY", "2"))

CORS_ORIGINS = [
    os.getenv("FRANKLIN_WEB_ORIGIN", "http://localhost:5173"),
    os.getenv("FRANKLIN_WEB_ORIGIN_ALT", "http://127.0.0.1:5173"),
    "http://localhost:8080",
    "http://localhost:8081",
    "http://localhost:8082",
    "http://127.0.0.1:8080",
    "http://127.0.0.1:8081",
    "http://127.0.0.1:8082",
]

app = FastAPI(title=APP_NAME)

app.add_middleware(
    CORSMiddleware,
    allow_origins=[o for o in CORS_ORIGINS if o],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ---------------- MODELS ----------------
class User(SQLModel, table=True):
    id: Optional[int] = Field(default=None, primary_key=True)
    email: str
    role: str  # client | contractor | admin
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))


class BidRequest(SQLModel, table=True):
    id: Optional[int] = Field(default=None, primary_key=True)
    client_id: int
    title: str
    description: str
    open: bool = True
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))


class Bid(SQLModel, table=True):
    id: Optional[int] = Field(default=None, primary_key=True)
    request_id: int
    contractor_id: int
    price: float
    message: Optional[str] = None
    accepted: bool = False
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))


class Contract(SQLModel, table=True):
    id: Optional[int] = Field(default=None, primary_key=True)
    bid_id: int
    contract_uuid: str
    status: str = "draft"
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))


class Audit(SQLModel, table=True):
    id: Optional[int] = Field(default=None, primary_key=True)
    event: str
    payload: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))


class CognitiveMemory(SQLModel, table=True):
    """Cognitive remembrance system with timestamps for PFS, Air Weaver, or Raspberry Pi"""
    id: Optional[int] = Field(default=None, primary_key=True)
    memory_key: str = Field(index=True)
    memory_value: str
    memory_type: str = Field(default="general")  # general | pfs | air_weaver | raspberry_pi
    context: Optional[str] = None
    embedding_vector: Optional[str] = None  # JSON string of vector for semantic search
    access_count: int = Field(default=0)
    last_accessed: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    expires_at: Optional[datetime] = None
    meta_data: Optional[str] = None  # JSON string for additional metadata (renamed from metadata)


class AIResponseCache(SQLModel, table=True):
    """Cache AI responses to reduce API calls"""
    id: Optional[int] = Field(default=None, primary_key=True)
    cache_key: str = Field(index=True, unique=True)  # Hash of prompt + provider + model
    provider: str
    model: str
    prompt_hash: str = Field(index=True)
    response_content: str
    response_metadata: Optional[str] = None
    hit_count: int = Field(default=0)
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    last_hit: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    expires_at: datetime  # TTL for cache


class WorkflowExecution(SQLModel, table=True):
    """Track workflow executions"""
    id: Optional[int] = Field(default=None, primary_key=True)
    workflow_id: str
    workflow_name: str
    input_data: str  # JSON
    output_data: Optional[str] = None  # JSON
    status: str = "pending"  # pending | running | completed | failed
    container_id: Optional[str] = None  # For containerized workflows
    error: Optional[str] = None
    started_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    completed_at: Optional[datetime] = None


class UploadedFile(SQLModel, table=True):
    """Track uploaded files for pipeline processing"""
    id: Optional[int] = Field(default=None, primary_key=True)
    file_uuid: str = Field(index=True, unique=True)
    filename: str
    file_path: str
    file_size: int
    file_type: str
    uploaded_by: Optional[int] = None  # user_id
    pipeline_id: Optional[str] = None
    processed: bool = Field(default=False)
    processing_result: Optional[str] = None
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))


SQLModel.metadata.create_all(engine)

# ---------------- HELPERS ----------------
def _now_ms() -> int:
    return int(datetime.now(timezone.utc).timestamp() * 1000)


def audit(event: str, payload: dict):
    with Session(engine) as s:
        s.add(Audit(event=event, payload=json.dumps(payload)))
        s.commit()


def create_token(user: User):
    return jwt.encode(
        {"sub": user.id, "role": user.role, "exp": datetime.now(timezone.utc) + timedelta(days=30)},
        SECRET,
        algorithm=JWT_ALGO,
    )


def get_user(req: Request) -> User:
    token = req.query_params.get("token")
    if not token:
        raise HTTPException(401, "Token required")
    try:
        data = jwt.decode(token, SECRET, algorithms=[JWT_ALGO])
    except Exception:
        raise HTTPException(401, "Invalid token") from None

    with Session(engine) as s:
        user = s.get(User, data["sub"])
        if not user:
            raise HTTPException(401, "User not found")
        return user


def generate_cache_key(prompt: str, provider: str, model: str) -> str:
    """Generate a unique cache key for AI requests"""
    import hashlib
    content = f"{provider}:{model}:{prompt}"
    return hashlib.sha256(content.encode()).hexdigest()


def get_cached_response(prompt: str, provider: str, model: str) -> Optional[str]:
    """Retrieve cached AI response if available and not expired"""
    cache_key = generate_cache_key(prompt, provider, model)
    with Session(engine) as s:
        cache_entry = s.exec(select(AIResponseCache).where(AIResponseCache.cache_key == cache_key)).first()
        if cache_entry and cache_entry.expires_at > datetime.now(timezone.utc):
            # Update hit count and last hit time
            cache_entry.hit_count += 1
            cache_entry.last_hit = datetime.now(timezone.utc)
            s.add(cache_entry)
            s.commit()
            return cache_entry.response_content
    return None


def cache_response(prompt: str, provider: str, model: str, response: str, ttl_hours: int = 24):
    """Cache AI response with TTL"""
    cache_key = generate_cache_key(prompt, provider, model)
    import hashlib
    prompt_hash = hashlib.sha256(prompt.encode()).hexdigest()
    
    with Session(engine) as s:
        # Check if entry exists
        existing = s.exec(select(AIResponseCache).where(AIResponseCache.cache_key == cache_key)).first()
        if existing:
            # Update existing
            existing.response_content = response
            existing.hit_count = 0
            existing.last_hit = datetime.now(timezone.utc)
            existing.expires_at = datetime.now(timezone.utc) + timedelta(hours=ttl_hours)
            s.add(existing)
        else:
            # Create new
            cache_entry = AIResponseCache(
                cache_key=cache_key,
                provider=provider,
                model=model,
                prompt_hash=prompt_hash,
                response_content=response,
                expires_at=datetime.now(timezone.utc) + timedelta(hours=ttl_hours)
            )
            s.add(cache_entry)
        s.commit()


def store_memory(key: str, value: str, memory_type: str = "general", context: Optional[str] = None, 
                 metadata: Optional[dict] = None, ttl_days: Optional[int] = None):
    """Store cognitive memory with timestamp"""
    with Session(engine) as s:
        expires_at = datetime.now(timezone.utc) + timedelta(days=ttl_days) if ttl_days else None
        memory = CognitiveMemory(
            memory_key=key,
            memory_value=value,
            memory_type=memory_type,
            context=context,
            meta_data=json.dumps(metadata) if metadata else None,
            expires_at=expires_at
        )
        s.add(memory)
        s.commit()
        s.refresh(memory)
        return memory


def retrieve_memory(key: str, memory_type: Optional[str] = None) -> Optional[str]:
    """Retrieve cognitive memory and update access tracking"""
    with Session(engine) as s:
        query = select(CognitiveMemory).where(CognitiveMemory.memory_key == key)
        if memory_type:
            query = query.where(CognitiveMemory.memory_type == memory_type)
        
        memory = s.exec(query).first()
        if memory:
            # Check if expired
            if memory.expires_at and memory.expires_at < datetime.now(timezone.utc):
                return None
            # Update access tracking
            memory.access_count += 1
            memory.last_accessed = datetime.now(timezone.utc)
            s.add(memory)
            s.commit()
            return memory.memory_value
    return None


# ---------------- AI MODELS ----------------
AIType = Literal["text", "image", "audio", "code", "analysis", "vision"]


class AIRequestModel(BaseModel):
    id: str = PydanticField(default_factory=lambda: str(uuid.uuid4()))
    type: AIType
    prompt: str
    provider: Optional[str] = None
    model: Optional[str] = None
    parameters: Optional[Dict[str, Any]] = None
    context: Optional[List[Any]] = None
    stream: Optional[bool] = False


class AIResponseModel(BaseModel):
    id: str
    provider: str
    model: str
    type: str
    content: Any
    timestamp: int


class MultiAgentRequestModel(BaseModel):
    prompt: str
    agents: Optional[List[str]] = None
    requestType: AIType = "text"


class PipelineExecuteRequestModel(BaseModel):
    pipelineId: str
    input: Any
    context: Optional[List[Any]] = None


PIPELINES: Dict[str, Dict[str, Any]] = {
    "content-gen": {
        "id": "content-gen",
        "name": "Content Generation Pipeline",
        "parallel": False,
        "stages": [
            {"id": "ideation", "name": "Idea Generation", "provider": "anthropic", "prompt": "Generate creative ideas for: {input}"},
            {"id": "expansion", "name": "Content Expansion", "provider": "openai", "prompt": "Expand and elaborate on this idea: {input}"},
            {"id": "optimization", "name": "SEO Optimization", "provider": "google", "prompt": "Optimize this content for SEO: {input}"},
        ],
    },
    "code-gen": {
        "id": "code-gen",
        "name": "Code Generation Pipeline",
        "parallel": False,
        "stages": [
            {"id": "architecture", "name": "Architecture Design", "provider": "anthropic", "prompt": "Design the architecture for: {input}"},
            {"id": "implementation", "name": "Code Implementation", "provider": "openai", "model": "gpt-4o-mini", "prompt": "Implement this architecture in code: {input}"},
            {"id": "review", "name": "Code Review", "provider": "anthropic", "prompt": "Review this code for best practices and security: {input}"},
            {"id": "documentation", "name": "Documentation", "provider": "openai", "prompt": "Generate comprehensive documentation for: {input}"},
        ],
    },
    "analysis": {
        "id": "analysis",
        "name": "Deep Analysis Pipeline",
        "parallel": False,
        "stages": [
            {"id": "data-extraction", "name": "Data Extraction", "provider": "google", "prompt": "Extract key data points from: {input}"},
            {"id": "pattern-recognition", "name": "Pattern Recognition", "provider": "anthropic", "prompt": "Identify patterns and trends in: {input}"},
            {"id": "insights", "name": "Insight Generation", "provider": "openai", "prompt": "Generate actionable insights from: {input}"},
            {"id": "recommendations", "name": "Recommendations", "provider": "anthropic", "prompt": "Provide strategic recommendations based on: {input}"},
        ],
    },
    "creative": {
        "id": "creative",
        "name": "Creative Generation Pipeline",
        "parallel": True,
        "stages": [
            {"id": "concept", "name": "Concept Development", "provider": "anthropic", "prompt": "Develop creative concepts for: {input}"},
            {"id": "visual", "name": "Visual Generation", "provider": "stability", "prompt": "Create visual representation: {input}"},
            {"id": "copy", "name": "Copy Writing", "provider": "openai", "prompt": "Write compelling copy for: {input}"},
        ],
    },
}


GLOBAL_SEMAPHORE = asyncio.Semaphore(max(1, GLOBAL_MAX_CONCURRENCY))
PROVIDER_SEMAPHORES: Dict[str, asyncio.Semaphore] = {
    "openai": asyncio.Semaphore(max(1, int(os.getenv("FRANKLIN_OPENAI_CONCURRENCY", str(DEFAULT_PROVIDER_CONCURRENCY))))),
    "anthropic": asyncio.Semaphore(max(1, int(os.getenv("FRANKLIN_ANTHROPIC_CONCURRENCY", str(DEFAULT_PROVIDER_CONCURRENCY))))),
    "google": asyncio.Semaphore(max(1, int(os.getenv("FRANKLIN_GOOGLE_CONCURRENCY", str(DEFAULT_PROVIDER_CONCURRENCY))))),
    "stability": asyncio.Semaphore(max(1, int(os.getenv("FRANKLIN_STABILITY_CONCURRENCY", str(DEFAULT_PROVIDER_CONCURRENCY))))),
}


async def _with_limits(provider: str, coro):
    async with GLOBAL_SEMAPHORE:
        sem = PROVIDER_SEMAPHORES.get(provider)
        if sem is None:
            return await coro
        async with sem:
            return await coro


async def _call_openai(req: AIRequestModel) -> AIResponseModel:
    model = req.model or "gpt-4o-mini"
    
    # Check cache first to reduce API calls
    cached = get_cached_response(req.prompt, "openai", model)
    if cached:
        print(f"Cache hit for OpenAI request {req.id}")
        return AIResponseModel(
            id=req.id,
            provider="openai",
            model=model,
            type=req.type,
            content=cached,
            timestamp=_now_ms(),
        )
    
    if not is_key_valid(OPENAI_API_KEY):
        print(f"Warning: OPENAI_API_KEY missing or invalid. Returning mock response for {req.id}")
        await asyncio.sleep(0.5)
        return AIResponseModel(
            id=req.id,
            provider="openai",
            model=model,
            type=req.type,
            content=f"[MOCK OPENAI] Processed: {req.prompt[:50]}...",
            timestamp=_now_ms(),
        )

    endpoint = "https://api.openai.com/v1/chat/completions"
    body: Dict[str, Any] = {
        "model": model,
        "messages": [
            {"role": "system", "content": "You are a sophisticated AI assistant."},
            *([(c) for c in (req.context or [])] if req.context else []),
            {"role": "user", "content": req.prompt},
        ],
        "temperature": (req.parameters or {}).get("temperature", 0.7),
        "max_tokens": (req.parameters or {}).get("maxTokens", 1200),
    }

    async def _do():
        async with httpx.AsyncClient(timeout=60) as client:
            resp = await client.post(
                endpoint,
                headers={"Authorization": f"Bearer {OPENAI_API_KEY}", "Content-Type": "application/json"},
                json=body,
            )
        if resp.status_code >= 400:
            raise HTTPException(resp.status_code, resp.text)
        data = resp.json()
        content = data["choices"][0]["message"]["content"]
        
        # Cache the response
        cache_response(req.prompt, "openai", model, content, ttl_hours=24)
        
        return AIResponseModel(
            id=req.id,
            provider="openai",
            model=model,
            type=req.type,
            content=content,
            timestamp=_now_ms(),
        )

    return await _with_limits("openai", _do())


async def _call_anthropic(req: AIRequestModel) -> AIResponseModel:
    model = req.model or "claude-3-5-sonnet-20241022"
    
    # Check cache first
    cached = get_cached_response(req.prompt, "anthropic", model)
    if cached:
        print(f"Cache hit for Anthropic request {req.id}")
        return AIResponseModel(
            id=req.id,
            provider="anthropic",
            model=model,
            type=req.type,
            content=cached,
            timestamp=_now_ms(),
        )
    
    if not is_key_valid(ANTHROPIC_API_KEY):
        print(f"Warning: ANTHROPIC_API_KEY missing or invalid. Returning mock response for {req.id}")
        await asyncio.sleep(0.5)
        return AIResponseModel(
            id=req.id,
            provider="anthropic",
            model=model,
            type=req.type,
            content=f"[MOCK ANTHROPIC] Processed: {req.prompt[:50]}...",
            timestamp=_now_ms(),
        )

    endpoint = "https://api.anthropic.com/v1/messages"
    body: Dict[str, Any] = {
        "model": model,
        "messages": [
            *([(c) for c in (req.context or [])] if req.context else []),
            {"role": "user", "content": req.prompt},
        ],
        "max_tokens": (req.parameters or {}).get("maxTokens", 1200),
        "temperature": (req.parameters or {}).get("temperature", 0.7),
    }

    async def _do():
        async with httpx.AsyncClient(timeout=60) as client:
            resp = await client.post(
                endpoint,
                headers={
                    "x-api-key": ANTHROPIC_API_KEY,
                    "anthropic-version": "2023-06-01",
                    "Content-Type": "application/json",
                },
                json=body,
            )
        if resp.status_code >= 400:
            raise HTTPException(resp.status_code, resp.text)
        data = resp.json()
        content = (data.get("content") or [{}])[0].get("text")
        
        # Cache the response
        cache_response(req.prompt, "anthropic", model, content, ttl_hours=24)
        
        return AIResponseModel(id=req.id, provider="anthropic", model=model, type=req.type, content=content, timestamp=_now_ms())

    return await _with_limits("anthropic", _do())


async def _call_google(req: AIRequestModel) -> AIResponseModel:
    model = req.model or "gemini-1.5-flash"
    
    # Check cache first
    cached = get_cached_response(req.prompt, "google", model)
    if cached:
        print(f"Cache hit for Google request {req.id}")
        return AIResponseModel(
            id=req.id,
            provider="google",
            model=model,
            type=req.type,
            content=cached,
            timestamp=_now_ms(),
        )
    
    if not is_key_valid(GOOGLE_API_KEY):
        print(f"Warning: GOOGLE_API_KEY missing or invalid. Returning mock response for {req.id}")
        await asyncio.sleep(0.5)
        return AIResponseModel(
            id=req.id,
            provider="google",
            model=model,
            type=req.type,
            content=f"[MOCK GOOGLE] Processed: {req.prompt[:50]}...",
            timestamp=_now_ms(),
        )

    endpoint = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={GOOGLE_API_KEY}"
    body: Dict[str, Any] = {
        "contents": [{"parts": [{"text": req.prompt}]}],
        "generationConfig": {
            "temperature": (req.parameters or {}).get("temperature", 0.7),
            "maxOutputTokens": (req.parameters or {}).get("maxTokens", 1200),
        },
    }

    async def _do():
        async with httpx.AsyncClient(timeout=60) as client:
            resp = await client.post(endpoint, headers={"Content-Type": "application/json"}, json=body)
        if resp.status_code >= 400:
            raise HTTPException(resp.status_code, resp.text)
        data = resp.json()
        content = data["candidates"][0]["content"]["parts"][0]["text"]
        
        # Cache the response
        cache_response(req.prompt, "google", model, content, ttl_hours=24)
        
        return AIResponseModel(id=req.id, provider="google", model=model, type=req.type, content=content, timestamp=_now_ms())

    return await _with_limits("google", _do())


async def _call_stability(req: AIRequestModel) -> AIResponseModel:
    if not is_key_valid(STABILITY_API_KEY):
        print(f"Warning: STABILITY_API_KEY missing or invalid. Returning mock response for {req.id}")
        await asyncio.sleep(0.5)
        # Return a 1x1 transparent pixel or similar placeholder
        return AIResponseModel(
            id=req.id,
            provider="stability",
            model="stable-diffusion-xl",
            type="image",
            content="data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNkYAAAAAYAAjCB0C8AAAAASUVORK5CYII=",
            timestamp=_now_ms(),
        )

    endpoint = "https://api.stability.ai/v1/generation/stable-diffusion-xl-1024-v1-0/text-to-image"
    body: Dict[str, Any] = {
        "text_prompts": [{"text": req.prompt, "weight": 1}],
        "cfg_scale": (req.parameters or {}).get("cfg_scale", 7),
        "height": (req.parameters or {}).get("height", 1024),
        "width": (req.parameters or {}).get("width", 1024),
        "steps": (req.parameters or {}).get("steps", 30),
        "samples": (req.parameters or {}).get("samples", 1),
    }

    async def _do():
        async with httpx.AsyncClient(timeout=120) as client:
            resp = await client.post(
                endpoint,
                headers={"Content-Type": "application/json", "Authorization": f"Bearer {STABILITY_API_KEY}"},
                json=body,
            )
        if resp.status_code >= 400:
            raise HTTPException(resp.status_code, resp.text)
        data = resp.json()
        b64 = data["artifacts"][0]["base64"]
        return AIResponseModel(
            id=req.id,
            provider="stability",
            model="stable-diffusion-xl",
            type="image",
            content=f"data:image/png;base64,{b64}",
            timestamp=_now_ms(),
        )

    return await _with_limits("stability", _do())


async def _execute_ai(req: AIRequestModel) -> AIResponseModel:
    provider = req.provider or "openai"
    if provider == "openai":
        return await _call_openai(req)
    if provider == "anthropic":
        return await _call_anthropic(req)
    if provider == "google":
        return await _call_google(req)
    if provider == "stability":
        return await _call_stability(req)
    raise HTTPException(400, f"Unsupported provider: {provider}")


# ---------------- ASYNC TASK QUEUE ----------------
TaskStatus = Literal["pending", "processing", "completed", "failed"]
TaskType = Literal["pipeline", "multi-agent"]


class TaskModel(BaseModel):
    id: str = PydanticField(default_factory=lambda: str(uuid.uuid4()))
    type: TaskType
    status: TaskStatus = "pending"
    request: Dict[str, Any]
    response: Optional[Any] = None
    error: Optional[str] = None
    startTime: Optional[int] = None
    endTime: Optional[int] = None


TASKS: Dict[str, TaskModel] = {}
QUEUE: Deque[str] = deque()
QUEUE_LOCK = asyncio.Lock()


async def _run_pipeline(pipeline_id: str, input_value: Any, context: Optional[List[Any]]):
    pipe = PIPELINES.get(pipeline_id)
    if not pipe:
        raise HTTPException(404, f"Pipeline {pipeline_id} not found")

    stages = pipe["stages"]
    parallel = bool(pipe.get("parallel"))

    def stage_prompt(stage: Dict[str, Any], value: Any) -> str:
        text = value if isinstance(value, str) else json.dumps(value)
        return (stage.get("prompt") or "{input}").replace("{input}", text)

    async def run_stage(stage: Dict[str, Any], value: Any):
        req_type: AIType = "image" if stage.get("provider") == "stability" else "text"
        req = AIRequestModel(
            type=req_type,
            prompt=stage_prompt(stage, value),
            provider=stage.get("provider"),
            model=stage.get("model"),
            context=context,
        )
        resp = await _execute_ai(req)
        return {"stage": stage["name"], "output": resp.content, "timestamp": _now_ms()}

    if parallel:
        results = await asyncio.gather(*(run_stage(s, input_value) for s in stages))
        return {"pipeline": pipe["name"], "stages": results, "timestamp": _now_ms()}

    current = input_value
    results: List[Dict[str, Any]] = []
    for s in stages:
        out = await run_stage(s, current)
        current = out["output"]
        results.append(out)

    return {"pipeline": pipe["name"], "stages": results, "finalOutput": current, "timestamp": _now_ms()}


async def _run_multi_agent(prompt: str, agents: List[str], request_type: AIType):
    async def call(agent: str):
        try:
            resp = await _execute_ai(AIRequestModel(type=request_type, prompt=prompt, provider=agent))
            return {"agent": agent, "status": "fulfilled", "response": resp.model_dump(), "error": None}
        except Exception as ex:
            return {"agent": agent, "status": "rejected", "response": None, "error": str(ex)}

    results = await asyncio.gather(*(call(a) for a in agents))
    fulfilled = [r for r in results if r["status"] == "fulfilled"]

    if fulfilled and ANTHROPIC_API_KEY:
        agg_prompt = f"Aggregate and synthesize these AI responses into a single comprehensive answer: {json.dumps(fulfilled)}"
        agg = await _execute_ai(AIRequestModel(type="text", prompt=agg_prompt, provider="anthropic"))
        return {"results": results, "aggregate": agg.model_dump()}

    return {"results": results}


async def _worker():
    while True:
        task_id: Optional[str] = None
        async with QUEUE_LOCK:
            if QUEUE:
                task_id = QUEUE.popleft()

        if not task_id:
            await asyncio.sleep(0.1)
            continue

        task = TASKS[task_id]
        task.status = "processing"
        task.startTime = _now_ms()
        TASKS[task_id] = task

        try:
            if task.type == "pipeline":
                task.response = await _run_pipeline(
                    pipeline_id=task.request["pipelineId"],
                    input_value=task.request["input"],
                    context=task.request.get("context"),
                )
            elif task.type == "multi-agent":
                task.response = await _run_multi_agent(
                    prompt=task.request["prompt"],
                    agents=task.request["agents"],
                    request_type=task.request.get("requestType", "text"),
                )
            else:
                raise HTTPException(400, f"Unknown task type: {task.type}")

            task.status = "completed"
        except Exception as ex:
            task.status = "failed"
            task.error = str(ex)
        finally:
            task.endTime = _now_ms()
            TASKS[task_id] = task


@app.on_event("startup")
async def _startup():
    asyncio.create_task(_worker())


if __name__ == "__main__":
    import uvicorn
    port = int(os.getenv("PORT", 8080))
    uvicorn.run(app, host="0.0.0.0", port=port)

# ---------------- FRONTEND ----------------
@app.get("/", response_class=HTMLResponse)
def home():
    return """
    <html>
    <head><title>Franklin OS</title></head>
    <body style="font-family:system-ui;background:#0f172a;color:#e5e7eb;padding:40px">
        <h1>Franklin OS • BidNova • Trinity</h1>
        <ul>
            <li>/health</li>
            <li>/api/ai/execute</li>
            <li>/api/ai/multi-agent</li>
            <li>/api/ai/pipelines</li>
            <li>/api/ai/pipelines/execute</li>
            <li>/api/orchestrator/tasks/&lt;taskId&gt;</li>
        </ul>
    </body>
    </html>
    """


# ---------------- AUTH ----------------
@app.post("/auth/register")
def register(email: str, role: str):
    with Session(engine) as s:
        user = s.exec(select(User).where(User.email == email)).first()
        if not user:
            user = User(email=email, role=role)
            s.add(user)
            s.commit()
            s.refresh(user)
            audit("user.register", {"email": email, "role": role})
        return {"token": create_token(user), "user_id": user.id}


# ---------------- BID REQUESTS ----------------
@app.post("/requests")
def create_request(title: str, description: str, req: Request):
    user = get_user(req)
    if user.role != "client":
        raise HTTPException(403, "Only clients can create requests")
    with Session(engine) as s:
        br = BidRequest(client_id=user.id, title=title, description=description)
        s.add(br)
        s.commit()
        s.refresh(br)
        audit("request.create", {"request_id": br.id})
        return br


@app.get("/requests")
def list_requests():
    with Session(engine) as s:
        return s.exec(select(BidRequest).where(BidRequest.open == True)).all()


# ---------------- BIDS ----------------
@app.post("/bids")
def submit_bid(request_id: int, price: float, message: str, req: Request):
    user = get_user(req)
    if user.role != "contractor":
        raise HTTPException(403, "Only contractors can bid")
    with Session(engine) as s:
        req_obj = s.get(BidRequest, request_id)
        if not req_obj or not req_obj.open:
            raise HTTPException(404, "Request closed or missing")
        bid = Bid(request_id=request_id, contractor_id=user.id, price=price, message=message)
        s.add(bid)
        s.commit()
        s.refresh(bid)
        audit("bid.submit", {"bid_id": bid.id})
        return bid


# ---------------- CONTRACTS ----------------
@app.post("/bids/{bid_id}/accept")
def accept_bid(bid_id: int, req: Request):
    user = get_user(req)
    with Session(engine) as s:
        bid = s.get(Bid, bid_id)
        if not bid:
            raise HTTPException(404, "Bid not found")
        br = s.get(BidRequest, bid.request_id)
        if br.client_id != user.id:
            raise HTTPException(403, "Not your request")
        bid.accepted = True
        br.open = False
        contract = Contract(bid_id=bid.id, contract_uuid=str(uuid.uuid4()))
        s.add(bid)
        s.add(br)
        s.add(contract)
        s.commit()
        s.refresh(contract)
        audit("contract.create", {"contract_uuid": contract.contract_uuid})
        return contract


# ---------------- ADMIN ----------------
@app.get("/admin/audit")
def read_audit():
    with Session(engine) as s:
        return s.exec(select(Audit).order_by(Audit.created_at.desc())).all()


# ---------------- HEALTH ----------------
@app.get("/health")
def health():
    return {"status": "ok", "system": APP_NAME, "time": datetime.now(timezone.utc).isoformat()}


# ---------------- AI ROUTES ----------------
@app.post("/api/ai/execute", response_model=AIResponseModel)
async def ai_execute(request: AIRequestModel):
    return await _execute_ai(request)


@app.get("/api/ai/pipelines")
def list_pipelines():
    return list(PIPELINES.values())


@app.post("/api/ai/pipelines/execute")
async def execute_pipeline(req: PipelineExecuteRequestModel):
    task = TaskModel(type="pipeline", request=req.model_dump())
    TASKS[task.id] = task
    async with QUEUE_LOCK:
        QUEUE.append(task.id)
    return {"taskId": task.id}


@app.post("/api/ai/multi-agent")
async def multi_agent(req: MultiAgentRequestModel):
    agents = req.agents or ["openai", "anthropic", "google", "stability"]
    task = TaskModel(type="multi-agent", request={"prompt": req.prompt, "agents": agents, "requestType": req.requestType})
    TASKS[task.id] = task
    async with QUEUE_LOCK:
        QUEUE.append(task.id)
    return {"taskId": task.id}


@app.get("/api/orchestrator/tasks/{task_id}")
def get_task(task_id: str):
    task = TASKS.get(task_id)
    if not task:
        raise HTTPException(404, "Task not found")
    return task.model_dump()


# ---------------- FILE UPLOAD FOR PIPELINES ----------------
from fastapi import UploadFile, File as FastAPIFile

@app.post("/api/files/upload")
async def upload_file(file: UploadFile = FastAPIFile(...)):
    """Upload file for pipeline processing"""
    import shutil
    
    # Create uploads directory if it doesn't exist
    uploads_dir = os.path.join(os.path.dirname(__file__), "uploads")
    os.makedirs(uploads_dir, exist_ok=True)
    
    # Generate unique file ID
    file_uuid = str(uuid.uuid4())
    file_ext = os.path.splitext(file.filename)[1]
    file_path = os.path.join(uploads_dir, f"{file_uuid}{file_ext}")
    
    # Save file
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    file_size = os.path.getsize(file_path)
    
    # Store in database
    with Session(engine) as s:
        uploaded_file = UploadedFile(
            file_uuid=file_uuid,
            filename=file.filename,
            file_path=file_path,
            file_size=file_size,
            file_type=file.content_type or "application/octet-stream"
        )
        s.add(uploaded_file)
        s.commit()
        s.refresh(uploaded_file)
        audit("file.upload", {"file_uuid": file_uuid, "filename": file.filename})
        
        return {
            "fileId": file_uuid,
            "filename": file.filename,
            "size": file_size,
            "type": file.content_type
        }


@app.get("/api/files/{file_id}")
def get_file_info(file_id: str):
    """Get uploaded file information"""
    with Session(engine) as s:
        file = s.exec(select(UploadedFile).where(UploadedFile.file_uuid == file_id)).first()
        if not file:
            raise HTTPException(404, "File not found")
        return file


# ---------------- COGNITIVE MEMORY ENDPOINTS ----------------
class MemoryStoreRequest(BaseModel):
    key: str
    value: str
    memory_type: str = "general"
    context: Optional[str] = None
    metadata: Optional[Dict[str, Any]] = None
    ttl_days: Optional[int] = None


@app.post("/api/memory/store")
def api_store_memory(req: MemoryStoreRequest):
    """Store cognitive memory"""
    memory = store_memory(
        key=req.key,
        value=req.value,
        memory_type=req.memory_type,
        context=req.context,
        metadata=req.metadata,
        ttl_days=req.ttl_days
    )
    audit("memory.store", {"key": req.key, "type": req.memory_type})
    return {"id": memory.id, "key": memory.memory_key, "status": "stored"}


@app.get("/api/memory/{key}")
def api_retrieve_memory(key: str, memory_type: Optional[str] = None):
    """Retrieve cognitive memory"""
    value = retrieve_memory(key, memory_type)
    if not value:
        raise HTTPException(404, "Memory not found or expired")
    return {"key": key, "value": value}


@app.get("/api/memory/list")
def list_memories(memory_type: Optional[str] = None, limit: int = 100):
    """List stored memories"""
    with Session(engine) as s:
        query = select(CognitiveMemory)
        if memory_type:
            query = query.where(CognitiveMemory.memory_type == memory_type)
        query = query.order_by(CognitiveMemory.created_at.desc()).limit(limit)
        memories = s.exec(query).all()
        return [
            {
                "id": m.id,
                "key": m.memory_key,
                "type": m.memory_type,
                "access_count": m.access_count,
                "created_at": m.created_at.isoformat(),
                "last_accessed": m.last_accessed.isoformat()
            }
            for m in memories
        ]


# ---------------- EXPORT ENDPOINTS ----------------
from fastapi.responses import StreamingResponse, FileResponse
import io

class ExportRequest(BaseModel):
    task_id: str
    format: Literal["excel", "word", "project", "audio", "jpeg"]


@app.post("/api/export")
async def export_data(req: ExportRequest):
    """Export task results to various formats"""
    task = TASKS.get(req.task_id)
    if not task:
        raise HTTPException(404, "Task not found")
    
    if task.status != "completed":
        raise HTTPException(400, "Task not completed yet")
    
    export_format = req.format
    
    # Excel Export
    if export_format == "excel":
        try:
            from openpyxl import Workbook
            wb = Workbook()
            ws = wb.active
            ws.title = "AI Results"
            
            # Add headers
            ws['A1'] = "Task ID"
            ws['A2'] = task.id
            ws['B1'] = "Type"
            ws['B2'] = task.type
            ws['C1'] = "Status"
            ws['C2'] = task.status
            
            # Add response data
            if task.response:
                ws['A4'] = "Results"
                ws['A5'] = json.dumps(task.response, indent=2)
            
            # Save to bytes
            output = io.BytesIO()
            wb.save(output)
            output.seek(0)
            
            return StreamingResponse(
                output,
                media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                headers={"Content-Disposition": f"attachment; filename=task_{task.id}.xlsx"}
            )
        except Exception as e:
            raise HTTPException(500, f"Excel export failed: {str(e)}")
    
    # Word Export
    elif export_format == "word":
        try:
            from docx import Document
            doc = Document()
            doc.add_heading(f"Task Results: {task.id}", 0)
            doc.add_paragraph(f"Type: {task.type}")
            doc.add_paragraph(f"Status: {task.status}")
            doc.add_paragraph("")
            doc.add_heading("Results", level=1)
            doc.add_paragraph(json.dumps(task.response, indent=2))
            
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            
            return StreamingResponse(
                output,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"attachment; filename=task_{task.id}.docx"}
            )
        except Exception as e:
            raise HTTPException(500, f"Word export failed: {str(e)}")
    
    # Project Export (JSON format for project management tools)
    elif export_format == "project":
        project_data = {
            "project_id": task.id,
            "task_type": task.type,
            "status": task.status,
            "created_at": task.request.get("timestamp", _now_ms()),
            "completed": task.endTime,
            "duration_ms": (task.endTime - task.startTime) if task.endTime and task.startTime else 0,
            "results": task.response
        }
        
        output = io.BytesIO(json.dumps(project_data, indent=2).encode())
        return StreamingResponse(
            output,
            media_type="application/json",
            headers={"Content-Disposition": f"attachment; filename=project_{task.id}.json"}
        )
    
    # Audio Export (Text-to-Speech)
    elif export_format == "audio":
        try:
            # Extract text content from task response
            text_content = ""
            if task.response:
                if isinstance(task.response, dict):
                    text_content = task.response.get("content", str(task.response))
                else:
                    text_content = str(task.response)
            
            # Use Google TTS if available
            if is_key_valid(GOOGLE_API_KEY):
                endpoint = "https://texttospeech.googleapis.com/v1/text:synthesize"
                body = {
                    "input": {"text": text_content[:5000]},  # Limit to 5000 chars
                    "voice": {"languageCode": "en-US", "name": "en-US-Neural2-C"},
                    "audioConfig": {"audioEncoding": "MP3"}
                }
                
                async with httpx.AsyncClient(timeout=60) as client:
                    resp = await client.post(
                        f"{endpoint}?key={GOOGLE_API_KEY}",
                        headers={"Content-Type": "application/json"},
                        json=body
                    )
                    
                if resp.status_code == 200:
                    import base64
                    audio_b64 = resp.json().get("audioContent")
                    audio_bytes = base64.b64decode(audio_b64)
                    
                    return StreamingResponse(
                        io.BytesIO(audio_bytes),
                        media_type="audio/mpeg",
                        headers={"Content-Disposition": f"attachment; filename=task_{task.id}.mp3"}
                    )
            
            # Fallback: return empty audio or error
            raise HTTPException(501, "Audio export requires Google API key")
        except Exception as e:
            raise HTTPException(500, f"Audio export failed: {str(e)}")
    
    # JPEG Export (for image results or visualization)
    elif export_format == "jpeg":
        try:
            from PIL import Image, ImageDraw, ImageFont
            
            # Create image with task info
            img = Image.new('RGB', (800, 600), color=(255, 255, 255))
            draw = ImageDraw.Draw(img)
            
            # Draw text
            text = f"Task: {task.id}\nType: {task.type}\nStatus: {task.status}"
            draw.text((20, 20), text, fill=(0, 0, 0))
            
            # If response contains image data, use that
            if task.response and isinstance(task.response, dict):
                if "content" in task.response and task.response.get("content", "").startswith("data:image"):
                    # Handle base64 image
                    import base64
                    img_data = task.response["content"].split(",")[1]
                    img_bytes = base64.b64decode(img_data)
                    img = Image.open(io.BytesIO(img_bytes))
            
            # Save to bytes
            output = io.BytesIO()
            img.save(output, format="JPEG", quality=95)
            output.seek(0)
            
            return StreamingResponse(
                output,
                media_type="image/jpeg",
                headers={"Content-Disposition": f"attachment; filename=task_{task.id}.jpg"}
            )
        except Exception as e:
            raise HTTPException(500, f"JPEG export failed: {str(e)}")
    
    raise HTTPException(400, f"Unsupported export format: {export_format}")


# ---------------- WORKFLOW TRACKING ----------------
@app.get("/api/workflows")
def list_workflows():
    """List workflow executions"""
    with Session(engine) as s:
        workflows = s.exec(select(WorkflowExecution).order_by(WorkflowExecution.started_at.desc()).limit(100)).all()
        return workflows


@app.get("/api/workflows/{workflow_id}")
def get_workflow(workflow_id: str):
    """Get workflow execution details"""
    with Session(engine) as s:
        workflow = s.exec(select(WorkflowExecution).where(WorkflowExecution.workflow_id == workflow_id)).first()
        if not workflow:
            raise HTTPException(404, "Workflow not found")
        return workflow
